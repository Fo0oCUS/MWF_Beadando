from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "MWF_Quiz_2026_dokumentacio.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                Path("C:/Windows/Fonts/arialbd.ttf"),
                Path("C:/Windows/Fonts/segoeuib.ttf"),
            ]
        )
    else:
        candidates.extend(
            [
                Path("C:/Windows/Fonts/arial.ttf"),
                Path("C:/Windows/Fonts/segoeui.ttf"),
            ]
        )

    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


FONT_TITLE = font(24, bold=True)
FONT_TEXT = font(18)
FONT_SMALL = font(16)


def ensure_assets() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, max_width: int, used_font) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""

    for word in words:
        trial = word if not current else current + " " + word
        width = draw.textbbox((0, 0), trial, font=used_font)[2]
        if width <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, lines: list[str], fill: tuple[int, int, int]) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=(60, 60, 60), width=3)
    draw.text((x1 + 18, y1 + 14), title, font=FONT_TITLE, fill=(20, 20, 20))

    y = y1 + 56
    for line in lines:
        draw.text((x1 + 18, y), line, font=FONT_TEXT, fill=(35, 35, 35))
        y += 28


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str | None = None) -> None:
    draw.line([start, end], fill=(40, 40, 40), width=4)

    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * sign, ey - 8), (ex - 16 * sign, ey + 8)]
    else:
        sign = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - 16 * sign), (ex + 8, ey - 16 * sign)]
    draw.polygon(points, fill=(40, 40, 40))

    if label:
        mx = (sx + ex) / 2
        my = (sy + ey) / 2
        bbox = draw.textbbox((0, 0), label, font=FONT_SMALL)
        padding = 8
        label_box = (
            mx - (bbox[2] - bbox[0]) / 2 - padding,
            my - (bbox[3] - bbox[1]) / 2 - padding,
            mx + (bbox[2] - bbox[0]) / 2 + padding,
            my + (bbox[3] - bbox[1]) / 2 + padding,
        )
        draw.rounded_rectangle(label_box, radius=8, fill=(255, 255, 255), outline=(180, 180, 180))
        draw.text((label_box[0] + padding, label_box[1] + padding), label, font=FONT_SMALL, fill=(40, 40, 40))


def centered_multiline(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, used_font) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 24
    lines = wrap_text(draw, text, max_width, used_font)
    heights = [draw.textbbox((0, 0), line, font=used_font)[3] for line in lines]
    total_height = sum(heights) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1 - total_height) / 2)

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=used_font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1 - width) / 2), y), line, font=used_font, fill=(20, 20, 20))
        y += (bbox[3] - bbox[1]) + 8


def component_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 950), (250, 248, 244))
    draw = ImageDraw.Draw(img)

    draw.text((40, 30), "UML komponensdiagram", font=font(28, bold=True), fill=(25, 25, 25))

    draw_box(
        draw,
        (60, 150, 340, 330),
        "Kliens",
        ["Bongeszo / frontend", "REST hivasok", "SignalR kapcsolat"],
        (232, 242, 252),
    )
    draw_box(
        draw,
        (470, 90, 850, 300),
        "QuizApp.WebApi",
        ["UsersController", "QuizController", "QuizHub", "JWT auth + exception kezeles"],
        (238, 248, 236),
    )
    draw_box(
        draw,
        (470, 370, 850, 620),
        "Quiz.DataAccess",
        ["IUserService / UserService", "IQuizService / QuizService", "QuizAppDbContext"],
        (255, 244, 214),
    )
    draw_box(
        draw,
        (980, 110, 1380, 310),
        "Shared",
        ["Request DTO-k", "Response DTO-k", "szerzodes a kliensekkel"],
        (245, 238, 251),
    )
    draw_box(
        draw,
        (980, 410, 1380, 640),
        "Adatbazis",
        ["SQL Server", "ASP.NET Identity tablakkal", "Quiz, Question, AppUser adatok"],
        (252, 236, 236),
    )

    arrow(draw, (340, 220), (470, 220), "REST / JSON")
    arrow(draw, (340, 270), (470, 500), "valos ideju es uzleti logika")
    arrow(draw, (850, 220), (980, 220), "DTO-k")
    arrow(draw, (660, 300), (660, 370), "szolgaltatasok")
    arrow(draw, (850, 520), (980, 520), "EF Core")

    footer = (
        "A kliens a web API vegpontjait hasznalja az adminisztracios es autentikacios muveletekhez, "
        "mig a SignalR kapcsolat a kerdescsere, a kviz lezarasa es a chat uzenetek azonnali tovabbitasat tamogatja."
    )
    centered_multiline(draw, (80, 720, 1520, 880), footer, FONT_TEXT)

    img.save(path)


def class_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 900), (248, 249, 245))
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "UML osztalydiagram", font=font(28, bold=True), fill=(25, 25, 25))

    draw_box(
        draw,
        (90, 170, 430, 500),
        "AppUser",
        ["Id: string", "Name: string", "Email: string", "RefreshToken: Guid?", "CreatedQuizzes: List<Quiz>"],
        (229, 241, 255),
    )
    draw_box(
        draw,
        (560, 130, 980, 610),
        "Quiz",
        [
            "Id: int",
            "Title: string",
            "UserId: string",
            "Questions: ICollection<Question>",
            "Messages: List<string>",
            "CurrentQuestionIndex: int",
            "IsPublished: bool",
            "JoinCode: string?",
            "Status: QuizStatus",
            "Players: List<string>?",
        ],
        (255, 246, 222),
    )
    draw_box(
        draw,
        (1110, 190, 1420, 520),
        "Question",
        ["Id: int", "QuizId: int", "Title: string", "Answers: List<string>", "CorrectAnswerIndex: int", "IsOpen: bool"],
        (235, 248, 233),
    )

    arrow(draw, (430, 320), (560, 320), "1 - *")
    arrow(draw, (980, 350), (1110, 350), "1 - *")

    legend_box = (100, 640, 1400, 820)
    draw.rounded_rectangle(legend_box, radius=18, fill=(255, 255, 255), outline=(160, 160, 160), width=2)
    centered_multiline(
        draw,
        legend_box,
        "A rendszer alapveto objektumai: a felhasznalo hozza letre a kvizeket, egy kviz pedig tobb kerdesbol all. "
        "A quiz osztaly tartalmazza a valos ideju jatek allapotat is, peldaul a csatlakozasi kodot, a jatekosokat es az uzeneteket.",
        FONT_TEXT,
    )

    img.save(path)


def schema_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 950), (249, 247, 243))
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "Adatbazis sema diagram", font=font(28, bold=True), fill=(25, 25, 25))

    draw_box(
        draw,
        (80, 140, 470, 530),
        "AspNetUsers",
        [
            "PK Id",
            "Name",
            "Email",
            "PasswordHash",
            "RefreshToken",
            "... tovabbi Identity mezok",
        ],
        (232, 242, 252),
    )
    draw_box(
        draw,
        (590, 120, 1030, 620),
        "Quizzes",
        [
            "PK Id",
            "FK UserId -> AspNetUsers.Id",
            "Title",
            "CurrentQuestionIndex",
            "IsPublished",
            "JoinCode",
            "Status",
            "Messages (szerializalt lista)",
            "Players (szerializalt lista)",
        ],
        (255, 244, 214),
    )
    draw_box(
        draw,
        (1150, 180, 1510, 570),
        "Questions",
        [
            "PK Id",
            "FK QuizId -> Quizzes.Id",
            "Title",
            "Answers (szerializalt lista)",
            "CorrectAnswerIndex",
            "IsOpen",
        ],
        (235, 248, 233),
    )

    arrow(draw, (470, 335), (590, 335), "1 - *")
    arrow(draw, (1030, 365), (1150, 365), "1 - *")

    identity_box = (80, 670, 1510, 860)
    draw.rounded_rectangle(identity_box, radius=18, fill=(255, 255, 255), outline=(160, 160, 160), width=2)
    centered_multiline(
        draw,
        identity_box,
        "A sajat alkalmazasi tablakat az ASP.NET Identity altal letrehozott szerep- es kapcsolotablak egeszitik ki "
        "(peldaul AspNetRoles, AspNetUserRoles, AspNetUserClaims). A projektben a kerdesvalaszok, chat uzenetek es jatekosnevek "
        "lista tipusu adatai szerializalt formaban kerulnek tarolasra.",
        FONT_TEXT,
    )

    img.save(path)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_doc_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(31, 56, 100)

    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run("MWF - Quiz 2026")
    run.bold = True
    run.font.size = Pt(20)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Webszolgáltatás dokumentáció").font.size = Pt(14)

    document.add_paragraph("")
    document.add_paragraph("")

    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.cell(0, 0).text = "Név"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "Neptun kód"
    table.cell(1, 1).text = ""
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_page_break()


def add_heading(document: Document, title: str, level: int = 1) -> None:
    document.add_heading(title, level=level)


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_scope_table(document: Document) -> None:
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        (
            "A feladat célja",
            "Egy .NET 8 alapú quiz webszolgáltatás elkészítése, amely támogatja a felhasználókezelést, a kvízek összeállítását és a játékosok csatlakozását.",
        ),
        (
            "Megvalósított alapfunkciók",
            "Regisztráció, bejelentkezés, saját kvízek létrehozása, lekérdezése, publikálása és a játékfolyamat kezelése.",
        ),
        (
            "Pontot érő kiegészítések",
            "SignalR alapú valós idejű működés, JWT + refresh token alapú authentikáció, szerkeszthető kvízek és élő chat.",
        ),
        (
            "A dokumentáció hatóköre",
            "A leírás kizárólag az alapfeladatra és a fent felsorolt megvalósított kiegészítésekre koncentrál, más választható feladatokra nem tér ki.",
        ),
    ]

    for row_index, (left, right) in enumerate(rows):
        left_cell = table.cell(row_index, 0)
        right_cell = table.cell(row_index, 1)
        left_cell.text = left
        right_cell.text = right
        set_cell_shading(left_cell, "DCE6F1")
        if left_cell.paragraphs and left_cell.paragraphs[0].runs:
            left_cell.paragraphs[0].runs[0].bold = True


def add_image(document: Document, path: Path, caption: str) -> None:
    document.add_picture(str(path), width=Inches(6.5))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cp = document.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(10)


def add_interface_table(document: Document) -> None:
    rows = [
        ("POST", "/users", "Felhasználó regisztráció", "nem"),
        ("POST", "/users/login", "Bejelentkezés, JWT és refresh token kiadása", "nem"),
        ("POST", "/users/logout", "Kijelentkezés", "igen"),
        ("POST", "/users/refresh", "Új hozzáférési token kérése refresh tokennel", "nem"),
        ("GET", "/users/{id}", "Saját felhasználói adat lekérése", "igen"),
        ("POST", "/quizzes", "Új kvíz létrehozása", "igen"),
        ("POST", "/quizzes/{id}/update", "Meglévő, még nem publikált kvíz szerkesztése", "igen"),
        ("GET", "/quizzes/{quizId}", "Kvíz lekérése tulajdonosként", "igen"),
        ("GET", "/quizzes/mine", "Saját kvízek listázása", "igen"),
        ("GET", "/quizzes/publish/{id}", "Kvíz publikálása, csatlakozási kód létrehozása", "igen"),
        ("POST", "/quizzes/join", "Csatlakozás közzétett kvízhez", "nem"),
        ("POST", "/quizzes/code", "Kvíz lekérése csatlakozási kód alapján", "nem"),
        ("GET", "/quizzes/{id}/next", "Következő kérdés indítása", "igen"),
        ("GET", "/quizzes/{id}/question/end", "Aktuális kérdés lezárása", "igen"),
        ("GET", "/quizzes/{id}/end", "Kvíz befejezése", "igen"),
        ("POST", "/quizzes/message", "Chat üzenet küldése a kvízhez", "nem"),
    ]

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    headers = ["Metódus", "Végpont", "Rövid leírás", "Hitelesítés"]
    for idx, value in enumerate(headers):
        header[idx].text = value
        set_cell_shading(header[idx], "D9EAF7")

    for method, endpoint, desc, auth in rows:
        cells = table.add_row().cells
        cells[0].text = method
        cells[1].text = endpoint
        cells[2].text = desc
        cells[3].text = auth


def add_signalr_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    headers = ["Elem", "Név", "Szerep"]
    for idx, value in enumerate(headers):
        header[idx].text = value
        set_cell_shading(header[idx], "E3F0D8")

    rows = [
        ("Hub", "/hubs/quiz", "A valós idejű kommunikáció belépési pontja"),
        ("Kliens metódus", "JoinQuizGroup(joinCode)", "Csatlakozás egy kvízhez tartozó csoporthoz"),
        ("Kliens metódus", "LeaveQuizGroup(joinCode)", "Kilépés a kvízcsoportból"),
        ("Szerver esemény", "QuestionChanged", "Értesítés új kérdés indításakor"),
        ("Szerver esemény", "QuestionClosed", "Értesítés kérdéslezáráskor, helyes válasz indexével"),
        ("Szerver esemény", "QuizEnded", "Értesítés a kvíz befejezéséről"),
        ("Szerver esemény", "MessageSent", "Új chat üzenet továbbítása"),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_test_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    headers = ["Tesztfájl", "Típus", "Lefedett működés"]
    for idx, value in enumerate(headers):
        header[idx].text = value
        set_cell_shading(header[idx], "FBE5D6")

    rows = [
        (
            "UserServiceTests",
            "egységteszt",
            "felhasználó létrehozása, bejelentkezés, aktuális felhasználó kezelése, kijelentkezés, jogosultságellenőrzés",
        ),
        (
            "QuizServiceTests",
            "egységteszt",
            "kvíz létrehozás, szerkesztés, publikálás, csatlakozás, kérdésléptetés, lezárás, üzenetküldés, befejezés",
        ),
        (
            "UsersControllerIntegrationTests",
            "integrációs teszt",
            "REST végpontok viselkedése felhasználókezelésnél, sikeres és hibás hitelesítési esetek",
        ),
        (
            "QuizControllerIntegrationTests",
            "integrációs teszt",
            "REST végpontok viselkedése kvízkezelésnél, jogosultsági és játékfolyamati esetek",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_testcase_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    headers = ["Teszt neve", "Előfeltétel", "Bemenet / művelet", "Elvárt eredmény"]
    for idx, value in enumerate(headers):
        header[idx].text = value
        set_cell_shading(header[idx], "FFF2CC")

    rows = [
        (
            "Sikeres bejelentkezés",
            "A felhasználó már regisztrált.",
            "POST /users/login helyes email és jelszó megadásával.",
            "A rendszer 200 OK választ ad vissza JWT tokennel, refresh tokennel és user azonosítóval.",
        ),
        (
            "Sikertelen bejelentkezés",
            "A felhasználó már regisztrált.",
            "POST /users/login hibás jelszóval.",
            "A rendszer hibaválaszt ad, és nem ad ki tokent.",
        ),
        (
            "Token megújítása",
            "A kliens rendelkezik érvényes refresh tokennel.",
            "POST /users/refresh a refresh token elküldésével.",
            "A rendszer új hozzáférési tokent ad vissza.",
        ),
        (
            "Kvíz létrehozása",
            "A felhasználó hitelesítve van.",
            "POST /quizzes érvényes címmel és kérdéslistával.",
            "A rendszer létrehozza a kvízt és 201 Created választ ad vissza.",
        ),
        (
            "Kvíz szerkesztése publikálás előtt",
            "A felhasználó a kvíz tulajdonosa, a kvíz még nincs publikálva.",
            "POST /quizzes/{id}/update módosított címmel vagy kérdésekkel.",
            "A rendszer elmenti a módosításokat és 202 Accepted választ ad vissza.",
        ),
        (
            "Csatlakozás közzétett kvízhez",
            "Létezik publikált kvíz érvényes join code-dal.",
            "POST /quizzes/join a join code és játékosnév megadásával.",
            "A játékos bekerül a kvíz résztvevői közé.",
        ),
        (
            "Következő kérdés indítása",
            "A kvíz publikált, a műveletet a tulajdonos végzi.",
            "GET /quizzes/{id}/next meghívása.",
            "Az aktuális kérdés indexe növekszik, és SignalR értesítés indul a kliensek felé.",
        ),
        (
            "Aktuális kérdés lezárása",
            "A kvízben már elindult egy kérdés.",
            "GET /quizzes/{id}/question/end meghívása.",
            "Az aktuális kérdés lezárul, a helyes válasz indexe továbbításra kerül a klienseknek.",
        ),
        (
            "Chat üzenet küldése",
            "A játékos már csatlakozott a kvízhez.",
            "POST /quizzes/message a join code, játékosnév és üzenet megadásával.",
            "Az üzenet elmentődik és valós időben megjelenik a többi kliens számára is.",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_document() -> None:
    ensure_assets()
    component_path = ASSETS / "component_diagram.png"
    class_path = ASSETS / "class_diagram.png"
    schema_path = ASSETS / "schema_diagram.png"

    component_diagram(component_path)
    class_diagram(class_path)
    schema_diagram(schema_path)

    document = Document()
    set_doc_style(document)
    add_title_page(document)

    add_heading(document, "1. Fejlesztő adatai")
    add_paragraph(document, "Név: ")
    add_paragraph(document, "Neptun kód: ")

    add_heading(document, "2. Feladatleírás")
    add_heading(document, "2.1. Feladat összefoglalása", level=2)
    add_paragraph(
        document,
        "A feladat egy valós idejű, Kahoot jellegű kvízrendszer elkészítése kliens–szerver architektúrában. "
        "A rendszerben a felhasználók saját kvízeket hozhatnak létre, a résztvevők pedig egy aktív kvízhez csatlakozva "
        "meghatározott időkereten belül válaszolhatnak a szerver által vezérelt kérdésekre.",
    )
    add_paragraph(
        document,
        "A webszolgáltatást REST architektúrát követő ASP.NET WebAPI alkalmazásként kell megvalósítani C# nyelven, "
        "MVC szemléletben. A rendszerhez egy résztvevői és egy adminisztrációs kliensfelület tartozik, amelyeknek "
        "együttműködve kell biztosítaniuk a teljes használatot és bemutathatóságot.",
    )
    add_bullet_list(
        document,
        [
            "A résztvevői felületen a felhasználók regisztrálhatnak, bejelentkezhetnek, illetve bejelentkezés nélkül is csatlakozhatnak egy aktív kvízhez egyedi azonosító megadásával.",
            "Az aktív kvíz során a kliens mindig az aktuális kérdést és annak válaszlehetőségeit jeleníti meg, egy kérdésre pontosan egy válasz adható le, kizárólag az aktív időszak alatt.",
            "A kérdés lezárása után a résztvevők megtekinthetik az adott kérdés eredményét.",
            "Az adminisztrációs felületen a bejelentkezett felhasználók megtekinthetik saját kvízeiket, új kvízt hozhatnak létre, valamint vezérelhetik a kvíz lebonyolítását.",
            "Egy kvíz létrehozásakor megadható a cím, a kérdések dinamikus száma, kérdésenként legalább két válaszlehetőség és a helyes válasz.",
            "A kérdések sorrendjét, aktiválását, lezárását és a következő kérdésre léptetést a szerver által vezérelt adminisztrációs folyamat szabályozza.",
            "A rendszernek biztosítania kell, hogy egy kvízben egyszerre tetszőleges számú résztvevő vehessen részt, és ugyanaz a felhasználó több kliensről is bejelentkezhessen.",
            "Az adatok relációs adatbázisban tárolódnak, kezelésük Entity Framework használatával történik, az authentikáció és authorizáció pedig ASP.NET Identity alapokra épül.",
            "A megoldásnak törekednie kell a biztonságos adatkezelésre, a felhasználóbarát működésre, a hibatűrésre, valamint a funkcionalitás integrációs tesztekkel történő ellenőrzésére.",
        ],
    )
    add_heading(document, "2.2. Választott feladatok", level=2)
    add_bullet_list(
        document,
        [
            "SignalR alapfeladat: a kérdések megjelenítése, lezárása és a kérdések közötti léptetés valós időben történik oldalfrissítés nélkül.",
            "OAuth2 alapú authentikáció: a hitelesítés refresh token flow mintával valósul meg, a hozzáférési tokenek JWT formátumúak.",
            "Szerkeszthető kvízek: az adminisztrációs felületen a saját, még nem elindított kvízek módosíthatók.",
            "Élő chat: a résztvevői felületen a kvízhez tartozó üzenetek valós időben jelennek meg, és a korábbi üzenetek is visszanézhetők.",
        ],
    )

    add_heading(document, "3. Funkcionális elemzés")
    add_bullet_list(
        document,
        [
            "A rendszer két fő szerepkört kezel: kvízkészítő felhasználó és csatlakozó játékos.",
            "A felhasználó regisztrálhat, majd bejelentkezés után saját kvízeket hozhat létre és listázhat.",
            "A kvíz címéből és kérdéseiből áll; egy kérdés szöveget, válaszlehetőségeket és helyes válasz indexet tartalmaz.",
            "A még nem publikált kvíz szerkeszthető, így a kérdések és a cím módosíthatók.",
            "A publikálás során a rendszer egy egyedi, hatjegyű csatlakozási kódot generál.",
            "A játékos a csatlakozási kód és egy választott név megadásával léphet be a kvízbe.",
            "A kvíz tulajdonosa léptetheti a kérdéseket, lezárhatja az aktuális kérdést, majd befejezheti a teljes játékot.",
            "A valós idejű működést a SignalR biztosítja, amely azonnal értesíti a klienseket a játékállapot változásairól.",
            "A chat funkció lehetővé teszi, hogy a csatlakozott játékosok üzenetet küldjenek az adott kvízhez.",
            "A hitelesítés JWT bearer tokennel történik, a hozzáférési token megújítása refresh tokennel valósul meg.",
        ],
    )

    add_heading(document, "4. Fejlesztői dokumentáció")
    add_heading(document, "4.1. Statikus terv", level=2)
    add_paragraph(
        document,
        "A megoldás négy fő projektre bontott: a webes API rétegre, az adatelérési és üzleti logikai rétegre, "
        "a közös DTO-kat tartalmazó projekt részre, valamint a tesztprojektre.",
    )
    add_image(document, component_path, "1. ábra - A rendszer komponensei és kapcsolatai")
    add_paragraph(
        document,
        "A webes réteg a kérések fogadásáért, a hitelesítésért, az útvonalkezelésért és a SignalR hub közzétételéért felel. "
        "Az üzleti szabályok a szolgáltatásokban, az állandó adatok kezelése pedig az Entity Framework Core alapú adatelérési rétegben található.",
    )
    add_image(document, class_path, "2. ábra - A legfontosabb alkalmazási osztályok")
    add_paragraph(
        document,
        "Az osztálydiagramon csak a lényegi domain elemek jelennek meg. A rendszer működését alapvetően az AppUser, Quiz és Question "
        "entitások írják le.",
    )

    add_heading(document, "4.2. Adatbázis felépítése", level=2)
    add_paragraph(
        document,
        "A projekt SQL Server adatbázist használ Entity Framework Core segítségével. A felhasználókezelés az ASP.NET Identity "
        "sémájára épül, amelyet az alkalmazás saját Quiz és Question táblái egészítenek ki.",
    )
    add_bullet_list(
        document,
        [
            "AspNetUsers: a felhasználók adatai, valamint a refresh token tárolása.",
            "Quizzes: a kvíz törzsadatai, állapota, csatlakozási kódja, játékoslistája és üzenetei.",
            "Questions: a kérdések szövege, válaszai, helyes válasza és nyitott vagy lezárt állapota.",
            "Identity táblák: a szerepkörök és kapcsolatok tárolására szolgáló ASP.NET Identity segédtáblák.",
        ],
    )
    add_image(document, schema_path, "3. ábra - A főbb adatbázistáblák kapcsolatai")
    add_paragraph(
        document,
        "A Quiz és a Question között egy-egy több kapcsolat található, ugyanígy egy felhasználó több kvíz létrehozója lehet. "
        "Az Answers, Messages és Players lista típusú adatok szerializált formában tárolódnak.",
    )

    add_heading(document, "4.3. A webszolgáltatás felülete", level=2)
    add_paragraph(
        document,
        "A rendszer REST alapú HTTP végpontokat és egy SignalR hubot tesz elérhetővé. A tulajdonosi műveletek hitelesítést igényelnek, "
        "a játékos oldali csatlakozási és chat műveletek viszont a join code alapján is használhatók.",
    )
    add_interface_table(document)
    add_paragraph(document, "")
    add_signalr_table(document)
    add_paragraph(
        document,
        "A JWT alapú hozzáférés rövid élettartamú bearer tokennel működik. A kliens a lejárt token helyett a "
        "POST /users/refresh végponton kérhet új hozzáférési tokent a korábban kapott refresh token felhasználásával.",
    )

    add_heading(document, "4.4. Tesztesetek", level=2)
    add_paragraph(
        document,
        "A tesztek a Quiz.Test projektben találhatók. A megoldás külön egységteszteket és integrációs teszteket tartalmaz.",
    )
    add_test_table(document)
    add_paragraph(
        document,
        "A tesztek lefedik a sikeres és hibás bejelentkezést, a jogosultságellenőrzést, a kvízek létrehozását és szerkesztését, "
        "a publikálási és csatlakozási folyamatot, a kérdésléptetést, a kérdéslezárást, a kvíz befejezését és a chat üzenetküldést.",
    )
    add_paragraph(
        document,
        "A legfontosabb funkcionális tesztesetek rövid leírása az alábbi táblázatban látható.",
    )
    add_testcase_table(document)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
